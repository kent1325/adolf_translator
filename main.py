"""
Document Translator - Streamlit App
User-friendly translation app for seniors
Supports drag-and-drop file upload and multiple translation engines
"""

import streamlit as st
from pathlib import Path
from typing import List, Optional
from docx import Document
import io
import re
from deep_translator import GoogleTranslator, MyMemoryTranslator, MicrosoftTranslator
import time
import subprocess
import tempfile
import os


class DocumentTranslator:
    def __init__(self, max_chunk_size: int = 1500):
        """Initialize the translator with conservative chunk size"""
        self.max_chunk_size = max_chunk_size
        self.languages = {
            "Auto-detect": "auto",
            "Chinese (Simplified)": "zh-CN",
            "Chinese (Traditional)": "zh-TW",
            "Spanish": "es",
            "Catalan": "ca",
            "English": "en",
            "French": "fr",
            "German": "de",
            "Italian": "it",
            "Portuguese": "pt",
            "Japanese": "ja",
            "Korean": "ko",
            "Arabic": "ar",
            "Russian": "ru",
            "Dutch": "nl",
            "Polish": "pl",
        }

        self.translators = {
            "Google Translate": "google",
            "Microsoft Translator": "microsoft",
            "MyMemory": "mymemory",
        }

        self.translator_descriptions = {
            "Google Translate": "🌐 Most popular, great for Chinese",
            "Microsoft Translator": "💼 Professional, good for business",
            "MyMemory": "📚 Community-driven, free alternative",
        }

    def convert_doc_to_docx(self, doc_file) -> io.BytesIO:
        """Convert .doc file to .docx format using LibreOffice"""
        try:
            # Create a temporary directory
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save uploaded .doc file
                doc_path = os.path.join(temp_dir, "input.doc")
                with open(doc_path, "wb") as f:
                    f.write(doc_file.read())

                # Convert using LibreOffice (soffice command)
                # This requires LibreOffice to be installed on the system
                result = subprocess.run(
                    [
                        "soffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        temp_dir,
                        doc_path,
                    ],
                    capture_output=True,
                    timeout=30,
                )

                # Read the converted file
                docx_path = os.path.join(temp_dir, "input.docx")
                if os.path.exists(docx_path):
                    with open(docx_path, "rb") as f:
                        docx_bytes = io.BytesIO(f.read())
                    return docx_bytes
                else:
                    raise Exception("Conversion failed - output file not created")

        except FileNotFoundError:
            raise Exception(
                "LibreOffice not found. Please install LibreOffice to convert .doc files."
            )
        except subprocess.TimeoutExpired:
            raise Exception("Conversion timed out. File may be too large or corrupted.")
        except Exception as e:
            raise Exception(f"Conversion error: {str(e)}")

    def extract_text_from_docx(self, file) -> str:
        """Extract text content from a Word document"""
        doc = Document(file)
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        return "\n\n".join(text)

    def smart_chunk_text(self, text: str) -> List[str]:
        """Split text into chunks that preserve context"""
        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = []
        current_size = 0

        for para in paragraphs:
            para_size = len(para)

            if para_size > self.max_chunk_size:
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_size = 0

                sentences = re.split(r"(?<=[.!?。！？])\s*", para)
                for sentence in sentences:
                    if len(sentence) > self.max_chunk_size:
                        for i in range(0, len(sentence), self.max_chunk_size):
                            chunks.append(sentence[i : i + self.max_chunk_size])
                    elif current_size + len(sentence) > self.max_chunk_size:
                        if current_chunk:
                            chunks.append("".join(current_chunk))
                        current_chunk = [sentence]
                        current_size = len(sentence)
                    else:
                        current_chunk.append(sentence)
                        current_size += len(sentence)

                if current_chunk:
                    chunks.append("".join(current_chunk))
                    current_chunk = []
                    current_size = 0

            elif current_size + para_size > self.max_chunk_size:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [para]
                current_size = para_size
            else:
                current_chunk.append(para)
                current_size += para_size + 2

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks

    def get_translator_instance(
        self, translator_type: str, source_code: str, target_code: str
    ):
        """Get the appropriate translator instance"""
        if translator_type == "google":
            return GoogleTranslator(source=source_code, target=target_code)
        elif translator_type == "microsoft":
            return MicrosoftTranslator(source=source_code, target=target_code)
        elif translator_type == "mymemory":
            return MyMemoryTranslator(source=source_code, target=target_code)
        else:
            return GoogleTranslator(source=source_code, target=target_code)

    def translate_text(
        self,
        text: str,
        target_code: str,
        source_code: str = "auto",
        translator_type: str = "google",
    ) -> str:
        """Translate text with automatic chunking"""
        if not text.strip():
            return ""

        if len(text) <= self.max_chunk_size:
            translator = GoogleTranslator(source=source_code, target=target_code)
            time.sleep(0.5)
            return translator.translate(text)

        chunks = self.smart_chunk_text(text)
        translated_chunks = []
        translator = GoogleTranslator(source=source_code, target=target_code)

        progress_bar = st.progress(0)
        status_text = st.empty()

        for i, chunk in enumerate(chunks):
            status_text.text(f"Translating chunk {i+1} of {len(chunks)}...")
            retry_count = 0
            max_retries = 3

            while retry_count < max_retries:
                try:
                    translated = translator.translate(chunk)
                    translated_chunks.append(translated)
                    time.sleep(1.5)  # Increased delay to avoid rate limits
                    break
                except Exception as e:
                    retry_count += 1
                    if retry_count < max_retries:
                        status_text.text(
                            f"Retrying chunk {i+1}... (attempt {retry_count + 1})"
                        )
                        time.sleep(3)  # Wait longer before retry
                    else:
                        st.warning(
                            f"⚠️ Could not translate chunk {i+1} after {max_retries} attempts. Skipping..."
                        )
                        translated_chunks.append(
                            f"[Translation unavailable for this section]"
                        )

            progress_bar.progress((i + 1) / len(chunks))

        progress_bar.empty()
        status_text.empty()

        return "\n\n".join(translated_chunks)

    def save_as_docx(self, text: str) -> io.BytesIO:
        """Save text as Word document in memory"""
        doc = Document()
        for paragraph in text.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


def main():
    st.set_page_config(page_title="Document Translator", page_icon="🌍", layout="wide")

    # Custom CSS for senior-friendly design
    st.markdown(
        """
        <style>
        .main {
            padding: 2rem;
        }
        .stButton>button {
            width: 100%;
            height: 3.5rem;
            font-size: 1.3rem;
            font-weight: bold;
            border-radius: 10px;
            margin-top: 0.5rem;
        }
        .stDownloadButton>button {
            width: 100%;
            height: 3.5rem;
            font-size: 1.3rem;
            font-weight: bold;
            border-radius: 10px;
            margin-top: 0.5rem;
        }
        .stTextArea textarea {
            font-size: 1.2rem;
        }
        h1 {
            font-size: 3rem !important;
            margin-bottom: 1rem;
        }
        h2 {
            font-size: 2rem !important;
        }
        .stSelectbox label {
            font-size: 1.3rem !important;
            font-weight: bold;
        }
        .stFileUploader label {
            font-size: 1.3rem !important;
            font-weight: bold;
        }
        div[data-testid="column"] {
            display: flex;
            flex-direction: column;
        }
        </style>
    """,
        unsafe_allow_html=True,
    )

    # Initialize session state
    if "translated_text" not in st.session_state:
        st.session_state.translated_text = ""
    if "last_file_id" not in st.session_state:
        st.session_state.last_file_id = None
    if "last_translator" not in st.session_state:
        st.session_state.last_translator = None
    if "current_text" not in st.session_state:
        st.session_state.current_text = ""

    # Initialize translator
    translator = DocumentTranslator()

    # Header
    st.title("🌍 Document Translator")
    st.markdown("### Simple and easy translation for your documents")
    st.markdown("---")

    # Translator selection
    st.markdown("## 🔧 Choose Your Translator")
    selected_translator = st.radio(
        "Select the translation service you'd like to use:",
        options=list(translator.translators.keys()),
        format_func=lambda x: f"{x} - {translator.translator_descriptions[x]}",
        horizontal=False,
        key="translator_choice",
    )
    translator_type = translator.translators[selected_translator]

    st.markdown("---")

    # Language selection
    col1, col2 = st.columns(2)

    with col1:
        # Remove Auto-detect from source language options
        source_languages = {
            k: v for k, v in translator.languages.items() if k != "Auto-detect"
        }
        source_lang = st.selectbox(
            "📖 Translate FROM:",
            options=list(source_languages.keys()),
            index=list(source_languages.keys()).index(
                "Chinese (Simplified)"
            ),  # Default to Chinese
            key="source_lang",
        )

    with col2:
        # Filter out Auto-detect for target language
        target_languages = {
            k: v for k, v in translator.languages.items() if k != "Auto-detect"
        }
        target_lang = st.selectbox(
            "📝 Translate TO:",
            options=list(target_languages.keys()),
            index=list(target_languages.keys()).index("Spanish"),  # Default to Spanish
            key="target_lang",
        )

    st.markdown("---")

    # Input section
    st.markdown("## 📄 Input Your Document")

    # File uploader
    uploaded_file = st.file_uploader(
        "Drag and drop a Word document here, or click to browse",
        type=["docx", "doc"],
        help="Upload a .docx or .doc file to translate",
    )

    # Future feature: Text input (currently disabled)
    # st.markdown("**OR**")
    # input_text = st.text_area(
    #     "Type or paste your text here:",
    #     height=200,
    #     placeholder="Enter the text you want to translate...",
    #     key='text_input'
    # )

    # Determine input source and extract text
    text_to_translate = ""
    current_file_id = None

    if uploaded_file is not None:
        try:
            # Check if file is .doc and convert to .docx
            file_extension = uploaded_file.name.split(".")[-1].lower()

            if file_extension == "doc":
                st.info("🔄 Converting .doc file to .docx format...")
                try:
                    converted_file = translator.convert_doc_to_docx(uploaded_file)
                    text_to_translate = translator.extract_text_from_docx(
                        converted_file
                    )
                except Exception as conv_error:
                    st.error(f"❌ Error converting .doc file: {conv_error}")
                    st.info(
                        "💡 Try saving your file as .docx format in Word and upload again."
                    )
                    text_to_translate = ""
            else:
                # Handle .docx files directly
                text_to_translate = translator.extract_text_from_docx(uploaded_file)

            if text_to_translate:
                # Create unique identifier for the file
                current_file_id = f"{uploaded_file.name}_{uploaded_file.size}"
                st.session_state.current_text = text_to_translate
                st.success(
                    f"✅ File loaded successfully! ({len(text_to_translate)} characters)"
                )
        except Exception as e:
            st.error(f"❌ Error reading file: {e}")
    # elif input_text:  # Uncomment when text input is re-enabled
    #     text_to_translate = input_text

    # Auto-translate when file is uploaded/changed or translator is changed
    needs_translation = False

    if current_file_id and current_file_id != st.session_state.last_file_id:
        needs_translation = True
        st.session_state.last_file_id = current_file_id

    if (
        translator_type != st.session_state.last_translator
        and st.session_state.current_text
    ):
        needs_translation = True
        st.session_state.last_translator = translator_type
        text_to_translate = st.session_state.current_text

    if needs_translation and text_to_translate.strip():
        with st.spinner("🔄 Translating..."):
            try:
                source_code = translator.languages[source_lang]
                target_code = translator.languages[target_lang]

                # Show document stats
                st.info(f"📊 Document size: {len(text_to_translate):,} characters")

                st.session_state.translated_text = translator.translate_text(
                    text_to_translate, target_code, source_code, translator_type
                )
                st.success("✅ Translation complete!")
                st.rerun()  # Force UI update
            except Exception as e:
                st.error(f"❌ Translation error: {e}")
                st.info(
                    "💡 Try selecting a different translator or splitting your document into smaller files."
                )
                st.session_state.translated_text = ""

    # Output section
    if st.session_state.translated_text:
        st.markdown("---")
        st.markdown("## ✨ Translated Text")

        # Display translated text (read-only) - use unique key based on content
        st.text_area(
            "Translation result:",
            value=st.session_state.translated_text,
            height=300,
            key=f"output_display_{hash(st.session_state.translated_text)}",
            disabled=True,
        )

        # Action buttons
        col1, col2 = st.columns(2)

        with col1:
            # Create download button (styled as primary for consistency)
            docx_buffer = translator.save_as_docx(st.session_state.translated_text)
            st.download_button(
                label="💾 Save as Word Document",
                data=docx_buffer,
                file_name=f"translated_{target_lang.lower().replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True,
            )

        with col2:
            if st.button(
                "📋 Copy to Clipboard", type="secondary", use_container_width=True
            ):
                st.code(st.session_state.translated_text, language=None)
                st.info("💡 Select the text above and copy it (Ctrl+C or Cmd+C)")

    # Footer
    st.markdown("---")
    st.markdown(
        """
        <div style='text-align: center; color: #666; padding: 2rem;'>
            <p style='font-size: 1.1rem;'>
                💡 <b>Tip:</b> Upload a .docx or .doc document and try different translators to find the best translation!
            </p>
            <p style='font-size: 1rem;'>
                Powered by multiple translation services
            </p>
        </div>
    """,
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    main()
